from docx import Document

def parse_docx(docx_file, subject):
    document = Document(docx_file)
    topics = []
    current_topic = None
    current_section = None

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.lower().startswith("topic:"):
            if current_topic:
                topics.append(current_topic)
            topic_title = text.split(":", 1)[1].strip()
            current_topic = {
                "topic": topic_title,
                "subject": subject
            }
            current_section = None
            continue

        lower = text.lower()
        if "definition" in lower:
            current_section = "definitions"
            continue
        if "example" in lower:
            current_section = "examples"
            continue
        if "process" in lower:
            current_section = "processes"
            continue
        if "equation" in lower:
            current_section = "equations"
            continue
        if "formula" in lower:
            current_section = "formulas"
            continue
        if "question" in lower:
            current_section = "questions"
            continue
        if lower.startswith("table:"):
            table_title = text.split(":", 1)[1].strip()
            current_section = "table"
            continue

        # 🔷 safeguard here:
        if current_topic is None:
            current_topic = {
                "topic": "Untitled Topic",
                "subject": subject
            }

        if current_section:
            if current_section not in current_topic:
                current_topic[current_section] = []
            current_topic[current_section].append(text)
        else:
            if "content" not in current_topic:
                current_topic["content"] = []
            current_topic["content"].append(text)

    if current_topic:
        topics.append(current_topic)

    # 🪄 If no topics but table exists
    if not topics:
        topics.append({
            "topic": "Untitled Topic",
            "subject": subject
        })

    i = 0
    for tbl in document.tables:
        if i >= len(topics):
            topics.append({
                "topic": "Untitled Topic",
                "subject": subject
            })
        while i < len(topics)-1 and "tables" in topics[i]:
            i += 1

        headers = [cell.text.strip() for cell in tbl.rows[0].cells]
        rows = []
        for row in tbl.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])

        # find table title if present in docx
        table_title = None
        for para in document.paragraphs:
            if para.text.strip().lower().startswith("table:"):
                table_title = para.text.strip().split(":", 1)[1].strip()
                break
        if not table_title:
            table_title = "Untitled Table"

        if "tables" not in topics[i]:
            topics[i]["tables"] = []
        topics[i]["tables"].append({
            "title": table_title,
            "headers": headers,
            "rows": rows
        })

    return topics
